# import necessary modules 
import sys
from urllib import request
import plotly.express as px
import plotly.io as pio
import pandas as pd
import json
import random
import docx

# call API, convert json into python list, then pick 5 elements at random from list
def get_park_list():
    response = request.urlopen("https://mn-state-parks.herokuapp.com/api/list").read()
    park_list = random.sample(json.loads(response), 5) # uses random.sample to get 5 unique parks at random from list

    # quit program if API cannot be reached, since program cannot function without succesful API call
    if response == None:
        sys.exit("Failed to reach API, exiting program")
        
    return park_list

# take list of parks and iterate through, doing an API call for each park, and putting detailed info in its own list
def get_detailed_info():
    park_info = []

    # do api call for each park in park list, append responses to new park_info list
    for park in get_park_list():
        response = request.urlopen(f"https://mn-state-parks.herokuapp.com/api/{park['park_id']}").read()
        park_info.append(json.loads(response))
    
    # if all API calls fail, exit program
    if park_info == []:
        sys.exit("Failed to reach API, exiting program")
    
    # if at least one successful API call, continue program. even if 4 fail, program can still work with data for one park.
    print(f"Succesfully retrieved data for {len(park_info)} out of 5 park(s).")
    print("Making document...")
    return park_info # return list of detailed info dictionaries 

# function to save images to working directory
def save_images(images):
    counter = 0

    # iterate through list of image urls and save each to working directory
    for img in images:
        request.urlretrieve(img, f"{counter}.jpg") 
        counter += 1 # use counter to name files 0.jpg, 1.jpg, etc...

# make a dataframe with pandas for later use with plotly to generate a map
def make_dataframe(park_info):
    # initialize empty list, which will form basis of data frame 
    data = []

    # for each park, append its name and coordinates into empty list
    for park in park_info:
        park_data = []
        park_data.append(park["name"])
        park_data.append(park["location"]["latitude"])
        park_data.append(park["location"]["longitude"])
        
        # append that list into the dataframe list
        data.append(park_data)

    return data

# generate map using plotly.express
def create_map(park_info):
    # create dataframe using pd.DataFrame. call make dataframe fuction to provide requisite list
    df = pd.DataFrame(make_dataframe(park_info))
    
    # create scatter_geo plot, pass it the data frame, and have it use the 3 columns in df for lon, lat, and text
    fig = px.scatter_geo(df,
        lon = 2,
        lat = 1,
        text = 0,
    )
    
    # set map paramaters, scope usa so it displays state borders, and fitbounds='locations' to automatically crop map
    fig.update_geos(
        scope = "usa", fitbounds = "locations",
        showcountries = True, showsubunits = True,
    )
    
    # add text to each marker on the map
    fig.update_traces(
        textposition="bottom center",
        mode="markers+text"
    )

    # save image to directory as map.png
    pio.write_image(fig, "map.png", scale=2, width=1000, height=800)


# function for making the word doc itself
def make_document():
    # initialize doc and add title, initialize park info list of dicts
    doc = docx.Document()
    doc.add_paragraph("Minnesota State Park Travel Guide", "Title")
    park_info = get_detailed_info()

    # create map of parks and add map to the start of the document 
    create_map(park_info)
    doc.add_picture("map.png", width=docx.shared.Inches(6))

    # iterate thru list of park info dicts
    for park in park_info:
        # pass current park into save_images, will download all relevant images for specific park into pwd
        save_images(park["park_images"])

        # add subtitle and 1 picture to start of entry for each park
        doc.add_paragraph(f"{park['name']}", "Title")
        doc.add_picture("0.jpg", width=docx.shared.Inches(6)) # manually set width to width of page

        # create bulleted list of highlights for each park
        doc.add_paragraph("Highlights", "Heading 2")
        for n, highlight in enumerate(park['highlights']):
            doc.add_paragraph(f"{highlight}", "List Bullet")

        # create paragraph with heading for each description found in park dict
        for header, description in park["park_information"].items():
            doc.add_paragraph(f"{header}", "Heading 2")
            doc.add_paragraph(f"{description}")
        
        # create heading for photos section
        # figure out how many images were downloaded by checking len(park["park_images"])
        # use for n in range to add picture to doc. don't use 0 since 0.jpg is used at start of section
        doc.add_paragraph(f"Photographs of {park['name']}", "Heading 1")
        for n in range(1, len(park["park_images"])):
            doc.add_picture(f"{n}.jpg", width=docx.shared.Inches(6))

        # pull address and website from park dict, and give each its own heading, along with the entire section 
        doc.add_paragraph("Further Information", "Heading 1")
        doc.add_paragraph("Address", "Heading 3")
        doc.add_paragraph(f"{park['address']}")
        doc.add_paragraph("Website", "Heading 3")
        doc.add_paragraph(f"{park['url']}")

        # add page break so next park can start at the top of its own page
        doc.add_page_break()
    
    # save doc once loop is finished
    doc.save('state_park_guide.docx')
    print("Done!")

if __name__ == "__main__":
    make_document()